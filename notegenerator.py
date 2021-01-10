from rake_nltk import Rake
from docx import Document
from itertools import zip_longest
import re

def generate_notes(text, blank_percentage):
    """
    returns a string with blanks for keywords and no repeats

    Parameters:
    text : str
        the text to manipulate and return with blanks
    blank_percentage : int
        the max ratio of words to be blanked
    """
    words = rake_keywords(text)
    blanks = len(text.split())*blank_percentage//100
    for i in range(blanks):
        if i > len(words) - 1:
            break;
        blank_space = '_' * len(words[i])
        text = text.replace(words[i], blank_space, 1)
    return text

def generate_docx(filepath, blank_percentage):
    document = Document(filepath)
    text = ""
    for paragraph in document.paragraphs:
        text += " " + paragraph.text

    words = rake_keywords(text)
    print(words)

    blanks = len(text.split())*blank_percentage//100
    for i in range(blanks):
        if i > len(words) - 1:
            break;

        for paragraph in document.paragraphs:
            if words[i] in paragraph.text:
                blank_space = '_' * len(words[i])
                paragraph.text = paragraph.text.replace(words[i], blank_space)
                document.save(filepath)
                break;

def rake_keywords(text):
    # use rake to extract keywords from the text

    rake = Rake(max_length=1)
    rake.extract_keywords_from_text(text)
    keywords = rake.get_ranked_phrases()
    keywords = [x for x in keywords if not bool(re.match('[^A-Za-z0-9]+', x))]
    return keywords
